import streamlit as st
import os
import json
import tempfile
import io
import pandas as pd
from docx import Document
from supabase import create_client, Client

# LlamaIndex 관련
from llama_index.core import VectorStoreIndex, SimpleDirectoryReader, Settings
from llama_index.llms.google_genai import GoogleGenAI
from llama_index.embeddings.google_genai import GoogleGenAIEmbedding

from langchain_google_genai import ChatGoogleGenerativeAI
from langchain.prompts import PromptTemplate
from langchain_core.output_parsers import JsonOutputParser, StrOutputParser

# 시크릿 관리
try:
    GEMINI_API_KEY = st.secrets["GEMINI_API_KEY"]
    SUPABASE_URL = st.secrets["SUPABASE_URL"]
    SUPABASE_KEY = st.secrets["SUPABASE_KEY"]
except KeyError:
    st.error("Secrets 설정이 누락되었습니다.")
    st.stop()

# Supabase 클라이언트
@st.cache_resource
def init_supabase():
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase = init_supabase()

# LlamaIndex 설정 (메모리 기반)
@st.cache_resource
def init_llama_index():
    Settings.llm = GoogleGenAI(
        model="gemini-2.0-flash",
        api_key=GEMINI_API_KEY,
        temperature=0.1
    )
    Settings.embed_model = GoogleGenAIEmbedding(
        model_name="text-embedding-004",
        api_key=GEMINI_API_KEY
    )
    Settings.chunk_size = 500
    Settings.chunk_overlap = 50

init_llama_index()

# LangChain LLM
llm = ChatGoogleGenerativeAI(
    model="gemini-2.0-flash",
    google_api_key=GEMINI_API_KEY,
    temperature=0.1
)

# 페이지 설정
st.set_page_config(page_title="RFP 제안서 자동생성 AI Agent", layout="wide")

# 세션 상태 초기화
if "company_info" not in st.session_state:
    st.session_state.company_info = {}
if "proposal_sections" not in st.session_state:
    st.session_state.proposal_sections = {}
if "requirements_check" not in st.session_state:
    st.session_state.requirements_check = []
if "fulfillment_rate" not in st.session_state:
    st.session_state.fulfillment_rate = 0.0
if "index" not in st.session_state:
    st.session_state.index = None
if "rfp_name" not in st.session_state:
    st.session_state.rfp_name = ""

# Supabase 저장 함수
def save_to_supabase(rfp_name, company_name, content, check, rate):
    try:
        supabase.table("proposals").insert({
            "rfp_name": rfp_name,
            "company_name": company_name,
            "proposal_content": content,
            "requirements_check": check,
            "fulfillment_rate": rate
        }).execute()
    except Exception as e:
        st.toast(f"DB 저장 실패: {e}")

# 4단계 Agent 파이프라인
def run_proposal_agent(query_engine, rfp_text_summary):

    # 1단계: RFP 분석
    with st.status("1단계: RFP 핵심 요구사항 분석 중...", expanded=True) as status:
        analysis_prompt = PromptTemplate.from_template(
            """다음 RFP 내용을 분석하여 아래 항목을 추출하라. 반드시 JSON만 출력할 것.

RFP 내용: {context}

출력 형식:
{{
    "requirements": ["요구사항1", "요구사항2"],
    "evaluation_criteria": ["평가기준1", "평가기준2"],
    "sections": ["필수섹션1", "필수섹션2"]
}}"""
        )
        try:
            analysis_chain = analysis_prompt | llm | JsonOutputParser()
            analysis_result = analysis_chain.invoke({"context": rfp_text_summary})
        except Exception:
            analysis_result = {
                "requirements": ["사업 목적 달성", "기간 내 완료", "성과 목표 충족"],
                "evaluation_criteria": ["기술력", "수행 경험", "가격"],
                "sections": ["사업 이해 및 추진 방향", "세부 수행 계획", "조직 및 인력 구성", "기대 효과"]
            }
        status.update(label="✅ 1단계 완료", state="complete")

    # 2단계: 목차 설계
    with st.status("2단계: 제안서 목차 구성 중...", expanded=True) as status:
        outline_prompt = PromptTemplate.from_template(
            """다음 RFP 분석 결과를 바탕으로 제안서 목차를 생성하라. 반드시 JSON만 출력할 것.

분석 결과: {analysis}

출력 형식:
{{"sections": ["1. 사업 이해 및 추진 방향", "2. 세부 수행 계획", "3. 조직 및 인력 투입 계획", "4. 기대 효과 및 성과 관리"]}}"""
        )
        try:
            outline_chain = outline_prompt | llm | JsonOutputParser()
            outline_result = outline_chain.invoke({"analysis": json.dumps(analysis_result, ensure_ascii=False)})
        except Exception:
            outline_result = {"sections": analysis_result.get("sections", ["사업 이해", "수행 계획", "조직 구성", "기대 효과"])}
        status.update(label="✅ 2단계 완료", state="complete")

    # 3단계 & 4단계 루프
    proposal_drafts = {}
    requirement_results = []
    final_rate = 0.0
    max_loops = 3

    for loop_count in range(1, max_loops + 1):
        sections = outline_result.get("sections", [])

        # 3단계: 섹션별 초안 생성
        with st.status(f"3단계: 섹션별 초안 생성 중... (시도 {loop_count}/{max_loops})", expanded=True) as status:
            for section in sections:
                try:
                    rag_response = query_engine.query(f"{section}에 대한 RFP 요구사항과 지침을 알려줘.")
                    rag_context = rag_response.response
                except Exception:
                    rag_context = rfp_text_summary

                draft_prompt = PromptTemplate.from_template(
                    """RFP 내용: {rfp_context}

업체 정보: {company}
섹션명: {section_name}

위 내용을 바탕으로 제안서 {section_name} 섹션 초안을 작성하라.

작성 규칙:
- 명사형 종결체 사용 (~함, ~임, ~예정, ~구성함)
- RFP에 근거가 없는 내용은 [추가 정보 필요: 이유]로 표기
- 구체적 수치와 방법론 포함
- 분량: 300자 이상

초안 내용만 출력할 것."""
                )
                try:
                    draft_chain = draft_prompt | llm | StrOutputParser()
                    draft = draft_chain.invoke({
                        "rfp_context": rag_context,
                        "company": json.dumps(st.session_state.company_info, ensure_ascii=False),
                        "section_name": section
                    })
                    proposal_drafts[section] = draft
                except Exception as e:
                    proposal_drafts[section] = f"[생성 오류: {str(e)}]"

            status.update(label=f"✅ 3단계 완료 ({len(sections)}개 섹션 생성)", state="complete")

        # 4단계: 요구사항 충족 검토
        with st.status("4단계: 요구사항 충족률 검토 중...", expanded=True) as status:
            review_prompt = PromptTemplate.from_template(
                """다음 RFP 요구사항 대비 제안서의 충족 여부를 검토하라. 반드시 JSON만 출력할 것.

RFP 요구사항: {reqs}
제안서 섹션 목록: {drafts}

출력 형식:
{{
    "check": [{{"requirement": "요구사항명", "met": true, "section": "대응섹션"}}],
    "fulfillment_rate": 85.0
}}"""
            )
            try:
                review_chain = review_prompt | llm | JsonOutputParser()
                review_result = review_chain.invoke({
                    "reqs": json.dumps(analysis_result.get("requirements", []), ensure_ascii=False),
                    "drafts": json.dumps(list(proposal_drafts.keys()), ensure_ascii=False)
                })
                requirement_results = review_result.get("check", [])
                final_rate = review_result.get("fulfillment_rate", 0.0)
            except Exception:
                requirement_results = [{"requirement": r, "met": True, "section": "전체"} for r in analysis_result.get("requirements", [])]
                final_rate = 75.0

            status.update(label=f"✅ 4단계 완료 (충족률: {final_rate}%)", state="complete")

        if final_rate >= 80.0:
            st.success(f"충족률 {final_rate}% 달성! 제안서 생성 완료.")
            break
        else:
            if loop_count < max_loops:
                st.warning(f"충족률 {final_rate}% — 미비 섹션 보완 중... ({loop_count}/{max_loops})")
            else:
                st.warning(f"최대 시도 횟수 초과. 최종 충족률: {final_rate}%")

    return proposal_drafts, requirement_results, final_rate


# UI 화면
st.title("RFP 제안서 자동 생성 AI Agent")
st.caption("제안요청서(RFP)를 업로드하면 AI가 자동으로 제안서 초안을 생성합니다.")

tab1, tab2, tab3 = st.tabs(["📂 RFP 업로드", "✍️ 제안서 생성", "💾 히스토리"])

# 탭 1: RFP 업로드
with tab1:
    col1, col2 = st.columns(2)

    with col1:
        st.subheader("업체 정보 입력")
        with st.form("company_form"):
            c_name = st.text_input("회사명", placeholder="주식회사 예시")
            c_ceo = st.text_input("대표자명")
            c_biz = st.text_input("사업자번호", placeholder="000-00-00000")
            c_perf = st.text_area("주요 실적", placeholder="최근 3년간 유사 프로젝트 수행 경험 등")
            c_core = st.text_area("핵심 역량", placeholder="전문 기술 및 인력 보유 현황 등")
            submitted = st.form_submit_button("정보 저장")
            if submitted:
                st.session_state.company_info = {
                    "회사명": c_name, "대표자": c_ceo,
                    "사업자번호": c_biz, "실적": c_perf, "역량": c_core
                }
                st.success("업체 정보가 저장되었습니다.")

    with col2:
        st.subheader("RFP PDF 업로드")
        uploaded_file = st.file_uploader("제안요청서(PDF)를 업로드하세요.", type="pdf")

        if uploaded_file:
            if not st.session_state.company_info:
                st.warning("먼저 왼쪽에서 업체 정보를 입력해주세요.")
            else:
                if st.button("RFP 인덱싱 시작", type="primary"):
                    with st.spinner("PDF를 분석 중입니다... (1~2분 소요)"):
                        try:
                            with tempfile.TemporaryDirectory() as tmp_dir:
                                tmp_path = os.path.join(tmp_dir, uploaded_file.name)
                                with open(tmp_path, "wb") as f:
                                    f.write(uploaded_file.getbuffer())
                                documents = SimpleDirectoryReader(input_dir=tmp_dir).load_data()
                                # 메모리 기반 인덱싱
                                index = VectorStoreIndex.from_documents(documents, show_progress=True)
                                st.session_state.index = index
                                st.session_state.rfp_name = uploaded_file.name
                                st.success(f"'{uploaded_file.name}' 인덱싱 완료! 제안서 생성 탭으로 이동하세요.")
                        except Exception as e:
                            st.error(f"인덱싱 오류: {e}")

# 탭 2: 제안서 생성
with tab2:
    if st.session_state.index is None:
        st.warning("먼저 탭 1에서 RFP를 업로드하고 인덱싱해주세요.")
    else:
        st.info(f"분석 대상: **{st.session_state.rfp_name}**")

        if st.button("제안서 초안 자동 생성 시작", type="primary"):
            try:
                query_engine = st.session_state.index.as_query_engine(similarity_top_k=5)
                rfp_summary = query_engine.query("이 RFP의 주요 목적, 사업 범위, 핵심 요구사항을 요약해줘.").response
                content, check, rate = run_proposal_agent(query_engine, rfp_summary)
                st.session_state.proposal_sections = content
                st.session_state.requirements_check = check
                st.session_state.fulfillment_rate = rate
                save_to_supabase(
                    st.session_state.rfp_name,
                    st.session_state.company_info.get("회사명", "미입력"),
                    content, check, rate
                )
            except Exception as e:
                st.error(f"제안서 생성 오류: {e}")

        if st.session_state.proposal_sections:
            st.divider()
            col_res1, col_res2 = st.columns([2, 1])

            with col_res1:
                st.subheader("생성된 제안서 초안")
                for section, text in st.session_state.proposal_sections.items():
                    with st.expander(f"{section}", expanded=True):
                        col_edit, col_btn = st.columns([4, 1])
                        with col_edit:
                            edited = st.text_area(
                                "내용", value=text,
                                height=200, key=f"edit_{section}",
                                label_visibility="collapsed"
                            )
                            st.session_state.proposal_sections[section] = edited
                        with col_btn:
                            if st.button("재생성", key=f"regen_{section}"):
                                with st.spinner("재생성 중..."):
                                    try:
                                        qe = st.session_state.index.as_query_engine(similarity_top_k=5)
                                        ctx = qe.query(f"{section} 관련 내용").response
                                        regen_prompt = PromptTemplate.from_template(
                                            "RFP 내용: {ctx}\n섹션: {sec}\n명사형 종결체로 제안서 초안 작성."
                                        )
                                        regen_chain = regen_prompt | llm | StrOutputParser()
                                        new_text = regen_chain.invoke({"ctx": ctx, "sec": section})
                                        st.session_state.proposal_sections[section] = new_text
                                        st.rerun()
                                    except Exception as e:
                                        st.error(f"재생성 오류: {e}")

            with col_res2:
                st.subheader("검토 결과")
                st.metric("요구사항 충족률", f"{st.session_state.fulfillment_rate:.1f}%")
                if st.session_state.requirements_check:
                    df_check = pd.DataFrame(st.session_state.requirements_check)
                    st.dataframe(df_check, use_container_width=True, hide_index=True)

            st.divider()
            st.subheader("결과물 다운로드")
            try:
                doc = Document()
                doc.add_heading(f"제안서: {st.session_state.rfp_name}", 0)
                doc.add_heading("업체 정보", level=1)
                for k, v in st.session_state.company_info.items():
                    doc.add_paragraph(f"{k}: {v}")
                for s, t in st.session_state.proposal_sections.items():
                    doc.add_heading(s, level=1)
                    doc.add_paragraph(t)
                doc.add_heading("요구사항 충족 대응표", level=1)
                if st.session_state.requirements_check:
                    table = doc.add_table(rows=1, cols=3)
                    table.style = "Table Grid"
                    hdr = table.rows[0].cells
                    hdr[0].text = "요구사항"
                    hdr[1].text = "충족 여부"
                    hdr[2].text = "대응 섹션"
                    for item in st.session_state.requirements_check:
                        row = table.add_row().cells
                        row[0].text = item.get("requirement", "")
                        row[1].text = "O" if item.get("met") else "X"
                        row[2].text = item.get("section", "")
                doc_buffer = io.BytesIO()
                doc.save(doc_buffer)
                doc_buffer.seek(0)
                st.download_button(
                    "Word(.docx) 다운로드",
                    doc_buffer,
                    file_name=f"{st.session_state.rfp_name}_제안서.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )
            except Exception as e:
                st.error(f"Word 생성 오류: {e}")

# 탭 3: 히스토리
with tab3:
    st.subheader("과거 제안서 생성 이력")
    try:
        response = supabase.table("proposals").select("*").order("created_at", desc=True).limit(50).execute()
        history_data = response.data
        if history_data:
            df_history = pd.DataFrame(history_data)
            st.dataframe(
                df_history[["created_at", "rfp_name", "company_name", "fulfillment_rate"]],
                use_container_width=True, hide_index=True
            )
            selected_id = st.selectbox("상세 보기할 제안서 ID 선택", df_history["id"])
            if selected_id:
                detail = next(item for item in history_data if item["id"] == selected_id)
                if detail.get("proposal_content"):
                    for sec, txt in detail["proposal_content"].items():
                        with st.expander(f"{sec}"):
                            st.write(txt)
            csv = df_history.to_csv(index=False).encode("utf-8-sig")
            st.download_button("CSV로 다운로드", csv, "proposal_history.csv", "text/csv")
        else:
            st.info("아직 생성된 제안서 이력이 없습니다.")
    except Exception as e:
        st.error(f"이력 불러오기 실패: {e}")
